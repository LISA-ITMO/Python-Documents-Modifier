from docx import Document
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from src.exceptions import NotSupportedFormat
from src.utils import Color, UnderlineStyle, FontStyle
from typing import List, Dict


class EDocx:
    """
    Class for edit DOCX document objects
    """
    def __init__(self, path: str) -> None:
        """
        :param path: path to DOCX document
        """
        if path[-5:] != '.docx':
            raise NotSupportedFormat(
                f'\'{path}\' should be .docx'
            )
        self.path: str = path
        self.document: Document = Document(path)
        self.autosave: bool = False

    def save(self, path=None) -> None:
        """
        Method for save changes in document
        """
        if path is None:
            self.document.save(self.path)
        else:
            self.document.save(path)

    @property
    def __para_ids(self) -> Dict[str, Paragraph]:
        res = {}
        key = '{http://schemas.microsoft.com/office/word/2010/wordml}paraId'
        for para in self.document.paragraphs:
            try:
                res[para.paragraph_format.element.attrib[key]] = para
            except KeyError as _ke:
                continue
        return res

    def add_comment_by_id(self, paraId: str, comment: str, author: str = 'EDocx') -> bool:
        """
        Add comment to paragraph by paraId
        :param paraId: paragraph id
        :param comment: a comment
        :param author: author's nickname
        :return: True - comment has been added, False - paragraph not found
        """
        para_ids = self.__para_ids
        if paraId not in para_ids:
            return False
        else:
            para_ids[paraId].add_comment(comment, author)
            if self.autosave:
                self.save()
            return True

    def all_para_attributes(self) -> List[Dict[str, str]]:
        """
        Get all attributes from all paragraphs
        :return: List of dicts (key: attribute, value: value)
        """
        attributes = []
        for para in self.document.paragraphs:
            attributes.append(para.paragraph_format.element.attrib)
        return attributes

    def edit_style_by_id(self, paraId: str,
                         size: int = None, color: Color = None,
                         fontStyle: FontStyle = None,
                         italic: bool = None, bold: bool = None,
                         underline: UnderlineStyle = None) -> bool:
        """
        Edit font style of paragraph by paraId
        :return: True - font style has been edited, False - paragraph not found
        """
        para_ids = self.__para_ids
        if paraId not in para_ids:
            return False
        for run in para_ids[paraId].runs:
            if size is not None:
                run.font.size = Pt(size)
            if color is not None:
                run.font.color.rgb = color
            if fontStyle is not None:
                run.font.name = fontStyle
            if italic is not None:
                run.font.italic = italic
            if bold is not None:
                run.font.bold = bold
            if underline is not None:
                run.font.underline = underline
        if self.autosave:
            self.save()
        return True
